import streamlit as st
import tempfile
import os
import re
import io
from datetime import datetime
import whisper
from streamlit_mic_recorder import mic_recorder

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle


# =========================
# CONFIGURACIÓN
# =========================

st.set_page_config(
    page_title="Sistema de Transcripción de Reuniones",
    page_icon="🎙️",
    layout="wide"
)

LIMITE_MINUTOS = 60
LIMITE_SEGUNDOS = LIMITE_MINUTOS * 60

FORMATOS_PERMITIDOS = ["wav", "mp3", "m4a", "ogg", "webm"]


# =========================
# CARGA MODELO
# =========================

@st.cache_resource
def cargar_modelo():
    return whisper.load_model("base")


modelo = cargar_modelo()


# =========================
# SESSION STATE
# =========================

if "transcripcion_total" not in st.session_state:
    st.session_state.transcripcion_total = ""


# =========================
# FUNCIONES
# =========================

def dividir_en_parrafos(texto):
    oraciones = re.split(r"(?<=[.!?])\s+", texto.strip())
    parrafos = []
    bloque = []

    for oracion in oraciones:
        if oracion:
            bloque.append(oracion.strip())

        if len(bloque) == 3:
            parrafos.append(" ".join(bloque))
            bloque = []

    if bloque:
        parrafos.append(" ".join(bloque))

    return parrafos


def buscar_tema(texto, tema):
    resultados = []

    if not texto or not tema:
        return resultados

    parrafos = dividir_en_parrafos(texto)

    for i, parrafo in enumerate(parrafos, start=1):
        if tema.lower() in parrafo.lower():
            resultados.append((i, parrafo))

    return resultados


def transcribir_audio(ruta_audio):
    resultado = modelo.transcribe(
        ruta_audio,
        language="es",
        fp16=False
    )
    return resultado["text"].strip()


def crear_word(nombre_asociacion, fecha_reunion, titulo_reunion, texto):
    documento = Document()

    section = documento.sections[0]
    section.top_margin = Pt(50)
    section.bottom_margin = Pt(50)
    section.left_margin = Pt(50)
    section.right_margin = Pt(50)

    titulo = documento.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("TRANSCRIPCIÓN DE REUNIÓN")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitulo = documento.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run(nombre_asociacion.upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(84, 130, 53)

    documento.add_paragraph("")

    p = documento.add_paragraph()
    p.add_run("Título de la reunión: ").bold = True
    p.add_run(titulo_reunion)

    p = documento.add_paragraph()
    p.add_run("Fecha de la reunión: ").bold = True
    p.add_run(str(fecha_reunion))

    p = documento.add_paragraph()
    p.add_run("Fecha de generación: ").bold = True
    p.add_run(datetime.now().strftime("%d/%m/%Y"))

    documento.add_paragraph("")

    encabezado = documento.add_paragraph()
    run = encabezado.add_run("Desarrollo de la reunión")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(31, 78, 121)

    parrafos = dividir_en_parrafos(texto)

    for parrafo in parrafos:
        p = documento.add_paragraph()
        p.style = "List Number"
        run = p.add_run(parrafo)
        run.font.size = Pt(11)

    documento.add_paragraph("")

    cierre = documento.add_paragraph()
    cierre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cierre.add_run("Documento generado automáticamente a partir de audio.")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    archivo = io.BytesIO()
    documento.save(archivo)
    archivo.seek(0)
    return archivo


def crear_pdf(nombre_asociacion, fecha_reunion, titulo_reunion, texto):
    archivo = io.BytesIO()

    doc = SimpleDocTemplate(
        archivo,
        pagesize=letter,
        rightMargin=50,
        leftMargin=50,
        topMargin=50,
        bottomMargin=50
    )

    estilos = getSampleStyleSheet()

    titulo_style = ParagraphStyle(
        "Titulo",
        parent=estilos["Title"],
        fontSize=18,
        textColor=colors.HexColor("#1F4E79"),
        alignment=1,
        spaceAfter=12
    )

    subtitulo_style = ParagraphStyle(
        "Subtitulo",
        parent=estilos["Heading2"],
        fontSize=13,
        textColor=colors.HexColor("#548235"),
        alignment=1,
        spaceAfter=18
    )

    encabezado_style = ParagraphStyle(
        "Encabezado",
        parent=estilos["Heading2"],
        fontSize=13,
        textColor=colors.HexColor("#1F4E79"),
        spaceBefore=15,
        spaceAfter=10
    )

    normal_style = ParagraphStyle(
        "Normal",
        parent=estilos["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=8
    )

    elementos = []

    elementos.append(Paragraph("TRANSCRIPCIÓN DE REUNIÓN", titulo_style))
    elementos.append(Paragraph(nombre_asociacion.upper(), subtitulo_style))

    elementos.append(Paragraph(f"<b>Título de la reunión:</b> {titulo_reunion}", normal_style))
    elementos.append(Paragraph(f"<b>Fecha de la reunión:</b> {fecha_reunion}", normal_style))
    elementos.append(Paragraph(f"<b>Fecha de generación:</b> {datetime.now().strftime('%d/%m/%Y')}", normal_style))

    elementos.append(Spacer(1, 12))
    elementos.append(Paragraph("Desarrollo de la reunión", encabezado_style))

    parrafos = dividir_en_parrafos(texto)

    for i, parrafo in enumerate(parrafos, start=1):
        elementos.append(Paragraph(f"<b>{i}.</b> {parrafo}", normal_style))

    elementos.append(Spacer(1, 20))
    elementos.append(Paragraph(
        "<i>Documento generado automáticamente a partir de audio.</i>",
        normal_style
    ))

    doc.build(elementos)
    archivo.seek(0)
    return archivo


def guardar_temporal_audio(audio_bytes, extension=".wav"):
    with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as temp_audio:
        temp_audio.write(audio_bytes)
        return temp_audio.name


# =========================
# INTERFAZ
# =========================

st.title("🎙️ Sistema de Transcripción y Gestión de Reuniones")
st.write(
    "Permite grabar reuniones, subir audios ya grabados, transcribirlos, buscar temas "
    "y descargar el resultado en Word, PDF o TXT."
)

with st.sidebar:
    st.header("📌 Datos del documento")

    nombre_asociacion = st.text_input(
        "Nombre de la asociación",
        value="Asociación de Desarrollo"
    )

    titulo_reunion = st.text_input(
        "Nombre o tema de la reunión",
        value="Reunión ordinaria"
    )

    fecha_reunion = st.date_input(
        "Fecha de la reunión",
        value=datetime.now()
    )

    st.info(f"Límite recomendado por audio: {LIMITE_MINUTOS} minutos.")


# =========================
# OPCIONES DE AUDIO
# =========================

st.subheader("🎧 Fuente de audio")

opcion_audio = st.radio(
    "Seleccione una opción:",
    [
        "🎙️ Grabar audio desde la app",
        "📂 Subir audio ya grabado"
    ],
    horizontal=True
)


# =========================
# GRABAR AUDIO
# =========================

if opcion_audio == "🎙️ Grabar audio desde la app":

    st.markdown("### 🎙️ Grabación directa")

    audio = mic_recorder(
        start_prompt="🎙️ Iniciar grabación",
        stop_prompt="⏹️ Detener grabación",
        just_once=False,
        use_container_width=True,
        key="grabador_audio"
    )

    if audio:
        audio_bytes = audio["bytes"]
        st.audio(audio_bytes, format="audio/wav")

        peso_mb = len(audio_bytes) / (1024 * 1024)
        st.info(f"Peso aproximado del audio grabado: {peso_mb:.2f} MB")

        if st.button("📝 Transcribir audio grabado"):
            temp_audio_path = None

            try:
                temp_audio_path = guardar_temporal_audio(audio_bytes, ".wav")

                with st.spinner("Transcribiendo audio grabado..."):
                    texto = transcribir_audio(temp_audio_path)

                if texto:
                    st.session_state.transcripcion_total += texto + "\n\n"
                    st.success("Audio transcrito correctamente.")
                else:
                    st.warning("No se detectó texto en el audio.")

            except Exception as e:
                st.error(f"Ocurrió un error al transcribir: {e}")

            finally:
                if temp_audio_path and os.path.exists(temp_audio_path):
                    os.remove(temp_audio_path)


# =========================
# SUBIR AUDIO
# =========================

if opcion_audio == "📂 Subir audio ya grabado":

    st.markdown("### 📂 Cargar archivo de audio")

    archivo_audio = st.file_uploader(
        "Suba un archivo de audio",
        type=FORMATOS_PERMITIDOS
    )

    if archivo_audio is not None:

        audio_bytes = archivo_audio.read()
        peso_mb = len(audio_bytes) / (1024 * 1024)

        st.audio(audio_bytes)
        st.info(f"Archivo cargado: {archivo_audio.name}")
        st.info(f"Peso del archivo: {peso_mb:.2f} MB")

        extension = os.path.splitext(archivo_audio.name)[1]

        if st.button("📝 Transcribir audio cargado"):
            temp_audio_path = None

            try:
                temp_audio_path = guardar_temporal_audio(audio_bytes, extension)

                with st.spinner("Transcribiendo audio cargado..."):
                    texto = transcribir_audio(temp_audio_path)

                if texto:
                    st.session_state.transcripcion_total += texto + "\n\n"
                    st.success("Audio cargado y transcrito correctamente.")
                else:
                    st.warning("No se detectó texto en el audio.")

            except Exception as e:
                st.error(f"Ocurrió un error al transcribir el archivo: {e}")

            finally:
                if temp_audio_path and os.path.exists(temp_audio_path):
                    os.remove(temp_audio_path)


# =========================
# TEXTO TRANSCRITO
# =========================

st.subheader("📄 Texto transcrito")

texto_editado = st.text_area(
    "Puede revisar o editar manualmente la transcripción:",
    value=st.session_state.transcripcion_total,
    height=350
)

st.session_state.transcripcion_total = texto_editado

col1, col2 = st.columns(2)

with col1:
    if st.button("🧹 Limpiar transcripción"):
        st.session_state.transcripcion_total = ""
        st.rerun()

with col2:
    st.download_button(
        label="⬇️ Descargar TXT",
        data=st.session_state.transcripcion_total,
        file_name="transcripcion_reunion.txt",
        mime="text/plain"
    )


# =========================
# BÚSQUEDA POR TEMA
# =========================

st.subheader("🔎 Búsqueda por tema")

tema_busqueda = st.text_input(
    "Escriba el tema a buscar. Ejemplo: contaminación del agua"
)

if st.button("Buscar tema"):
    resultados = buscar_tema(
        st.session_state.transcripcion_total,
        tema_busqueda
    )

    if resultados:
        st.success(f"Se encontraron {len(resultados)} coincidencias.")
        for numero, parrafo in resultados:
            st.markdown(f"**Punto / párrafo {numero}:**")
            st.info(parrafo)
    else:
        st.warning("No se encontraron menciones sobre ese tema.")


# =========================
# DESCARGAS
# =========================

st.subheader("📘 Generar documentos")

texto_para_documento = st.session_state.transcripcion_total

if texto_para_documento.strip():

    archivo_word = crear_word(
        nombre_asociacion=nombre_asociacion,
        fecha_reunion=fecha_reunion,
        titulo_reunion=titulo_reunion,
        texto=texto_para_documento
    )

    archivo_pdf = crear_pdf(
        nombre_asociacion=nombre_asociacion,
        fecha_reunion=fecha_reunion,
        titulo_reunion=titulo_reunion,
        texto=texto_para_documento
    )

    col1, col2 = st.columns(2)

    with col1:
        st.download_button(
            label="⬇️ Descargar Word",
            data=archivo_word,
            file_name="transcripcion_reunion.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    with col2:
        st.download_button(
            label="⬇️ Descargar PDF",
            data=archivo_pdf,
            file_name="transcripcion_reunion.pdf",
            mime="application/pdf"
        )

else:
    st.warning("Primero debe grabar, subir o escribir una transcripción para generar documentos.")


st.caption(
    "Sistema demo: grabación de audio, carga de audios existentes, transcripción, búsqueda temática y generación de Word/PDF."
)
